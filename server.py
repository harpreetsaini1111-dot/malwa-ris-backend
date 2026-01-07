import os
import io
import time
import html

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from google.cloud import storage

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# =========================================================
# APP
# =========================================================
app = Flask(__name__)
CORS(app)

# =========================================================
# GCS SETUP
# =========================================================
BUCKET_NAME = os.environ.get("GCS_BUCKET")
if not BUCKET_NAME:
    raise RuntimeError("GCS_BUCKET environment variable not set")

client = storage.Client()
bucket = client.bucket(BUCKET_NAME)

# =========================================================
# HELPERS
# =========================================================
def auto_name(path: str) -> str:
    blob = bucket.blob(path)
    if not blob.exists():
        return path
    base, ext = os.path.splitext(path)
    return f"{base}_{int(time.time())}{ext}"

def upload_bytes(path: str, data: bytes, content_type: str):
    path = auto_name(path)
    blob = bucket.blob(path)
    blob.upload_from_string(data, content_type=content_type)
    return path

def html_to_docx(html_text: str) -> bytes:
    doc = Document()
    clean = html_text.replace("<br>", "\n").replace("</p>", "\n").replace("<p>", "")
    for line in clean.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def html_to_pdf(html_text: str) -> bytes:
    bio = io.BytesIO()
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(bio)
    pdf.build([Paragraph(html.escape(html_text), styles["Normal"])])
    bio.seek(0)
    return bio.read()

def docx_to_html(file_obj) -> str:
    doc = Document(file_obj)
    blocks = []
    for p in doc.paragraphs:
        text = html.escape(p.text)
        if text.strip():
            blocks.append(f"<p>{text}</p>")
    return "\n".join(blocks)

# =========================================================
# BASIC
# =========================================================
@app.route("/")
def home():
    return "SERVER RUNNING", 200

# =========================================================
# SAVE REPORT (EDITOR → GCS DOCX)
# =========================================================
@app.route("/save_report", methods=["POST"])
def save_report():
    d = request.json
    modality = d["modality"]
    patient = d.get("patient_name", "UNKNOWN").replace(" ", "_")
    pid = d.get("patient_id", "AUTO")
    date = d.get("date", "")
    html_text = d.get("content", "")

    filename = f"{patient}_{pid}_{date}_{modality}.docx"
    path = f"reports/{modality}/{filename}"

    saved = upload_bytes(
        path,
        html_to_docx(html_text),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    return jsonify({"saved": [saved], "skipped": []})

# =========================================================
# DOWNLOAD FROM EDITOR (HTML → DOCX / PDF)
# =========================================================
@app.route("/download_from_editor", methods=["POST"])
def download_from_editor():
    d = request.json

    modality = d.get("modality", "REPORT")
    patient = d.get("patient_name", "UNKNOWN").replace(" ", "_")
    pid = d.get("patient_id", "AUTO")
    date = d.get("date", "")
    html_text = d.get("content", "")
    rtype = d.get("type", "docx")

    filename = f"{patient}_{pid}_{date}_{modality}.{rtype}"

    if rtype == "docx":
        return send_file(
            io.BytesIO(html_to_docx(html_text)),
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if rtype == "pdf":
        return send_file(
            io.BytesIO(html_to_pdf(html_text)),
            as_attachment=True,
            download_name=filename,
            mimetype="application/pdf"
        )

    return jsonify({"error": "Unsupported type"}), 400

# =========================================================
# SAVE TEMPLATE
# =========================================================
@app.route("/save_template", methods=["POST"])
def save_template():
    d = request.json
    modality = d["modality"]
    name = d["name"].replace(" ", "_") + ".html"

    path = f"templates/{modality}/{name}"
    saved = upload_bytes(path, d["content"].encode(), "text/html")
    return jsonify({"saved": [saved], "skipped": []})

# =========================================================
# LIST TEMPLATES
# =========================================================
@app.route("/list_templates")
def list_templates():
    modality = request.args["modality"]
    prefix = f"templates/{modality}/"

    templates = [
        b.name.split("/")[-1]
        for b in bucket.list_blobs(prefix=prefix)
        if not b.name.endswith("/")
    ]
    return jsonify({"templates": templates})

# =========================================================
# LOAD TEMPLATE
# =========================================================
@app.route("/load_template")
def load_template():
    modality = request.args["modality"]
    filename = request.args["filename"]

    blob = bucket.blob(f"templates/{modality}/{filename}")
    return jsonify({"content": blob.download_as_text()})

# =========================================================
# BATCH UPLOAD TEMPLATES
# =========================================================
@app.route("/batch_upload_templates", methods=["POST"])
def batch_upload_templates():
    modality = request.form["modality"]
    saved, skipped = [], []

    for f in request.files.getlist("files"):
        try:
            base = os.path.splitext(f.filename)[0]
            if f.filename.lower().endswith(".html"):
                saved.append(upload_bytes(
                    f"templates/{modality}/{base}.html",
                    f.read(),
                    "text/html"
                ))
            elif f.filename.lower().endswith(".docx"):
                html_content = docx_to_html(f)
                saved.append(upload_bytes(
                    f"templates/{modality}/{base}.html",
                    html_content.encode(),
                    "text/html"
                ))
        except Exception as e:
            skipped.append({"file": f.filename, "reason": str(e)})

    return jsonify({"saved": saved, "skipped": skipped})

# =========================================================
# BATCH UPLOAD REPORTS
# =========================================================
@app.route("/batch_upload_reports_auto", methods=["POST"])
def batch_upload_reports_auto():
    modality = request.form["modality"]
    saved, skipped = [], []

    for f in request.files.getlist("files"):
        try:
            saved.append(upload_bytes(
                f"reports/{modality}/{f.filename}",
                f.read(),
                f.content_type
            ))
        except Exception as e:
            skipped.append({"file": f.filename, "reason": str(e)})

    return jsonify({"saved": saved, "skipped": skipped})

# =========================================================
# LIST REPORTS
# =========================================================
@app.route("/list_reports")
def list_reports():
    modality = request.args.get("modality")
    prefix = f"reports/{modality}/" if modality else "reports/"

    reports = []
    for b in bucket.list_blobs(prefix=prefix):
        if not b.name.endswith("/"):
            reports.append({
                "filename": b.name.split("/")[-1],
                "path": b.name
            })

    return jsonify({"reports": reports})

# =========================================================
# DOWNLOAD SAVED REPORT
# =========================================================
@app.route("/download_report")
def download_report():
    path = request.args["path"]
    rtype = request.args.get("type", "docx")

    blob = bucket.blob(path)
    data = blob.download_as_bytes()

    if rtype == "pdf":
        return send_file(
            io.BytesIO(html_to_pdf(data.decode(errors="ignore"))),
            as_attachment=True,
            download_name=os.path.basename(path).replace(".docx", ".pdf"),
            mimetype="application/pdf"
        )

    return send_file(
        io.BytesIO(data),
        as_attachment=True,
        download_name=os.path.basename(path),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# =========================================================
# RUN
# =========================================================
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)






