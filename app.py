import io
import datetime
from flask import Flask, request, jsonify, Response, render_template
from flask_cors import CORS
from docx import Document

from storage import get_bucket   # your existing GCS helper

app = Flask(__name__)
CORS(app)   # 🔴 REQUIRED for browser → Render calls

# ---------------- HEALTH ----------------

@app.route("/health")
def health():
    return "OK"

# ---------------- HOME ----------------

@app.route("/")
def index():
    return "Malwa RIS Backend Running"

# ---------------- TEMPLATE APIs ----------------

@app.route("/templates", methods=["GET"])
def list_templates():
    modality = request.args.get("modality")
    bucket = get_bucket()

    blobs = bucket.list_blobs(prefix=f"templates/{modality}/")
    out = []

    for b in blobs:
        if b.name.endswith(".html"):
            out.append({
                "id": b.name,
                "name": b.name.split("/")[-1]
            })

    return jsonify(out)

@app.route("/templates/<path:blob_name>", methods=["GET"])
def load_template(blob_name):
    bucket = get_bucket()
    blob = bucket.blob(blob_name)

    return Response(
        blob.download_as_text(),
        mimetype="text/html"
    )

@app.route("/templates/save", methods=["POST"])
def save_template():
    data = request.json
    bucket = get_bucket()

    path = f"templates/{data['modality']}/{data['name']}"

    blob = bucket.blob(path)
    blob.upload_from_string(
        data["html"],
        content_type="text/html"
    )

    return jsonify({"status": "ok", "path": path})

# ---------------- REPORT API ----------------

@app.route("/reports/save", methods=["POST"])
def save_report():
    data = request.json
    patient = data["patient"]
    html = data["html"]

    # Create DOCX
    doc = Document()
    doc.add_heading("Radiology Report", level=1)

    doc.add_paragraph(
        f"Name: {patient.get('name','')}    "
        f"Age/Sex: {patient.get('age','')}/{patient.get('sex','')}    "
        f"UHID: {patient.get('uhid','')}\n"
        f"Ref: {patient.get('ref','')}    "
        f"Date: {patient.get('date','')}"
    )

    doc.add_paragraph("")
    doc.add_paragraph(html)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    today = datetime.date.today()
    filename = f"{patient.get('name','REPORT')}_{today}.docx"

    path = f"reports/{today.year}/{today.month}/{filename}"

    bucket = get_bucket()
    blob = bucket.blob(path)

    blob.upload_from_file(
        bio,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    return jsonify({
        "status": "ok",
        "path": path
    })

# ---------------- GCS TEST ----------------

@app.route("/test-gcs")
def test_gcs():
    bucket = get_bucket()
    blob = bucket.blob("test/hello.txt")
    blob.upload_from_string("GCS working")
    return jsonify({"status": "ok"})

# ---------------- MAIN ----------------

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)






